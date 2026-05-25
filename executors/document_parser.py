"""
executors/document_parser.py — REAL document parser v8
Algoritma 1: ekstrak teks dari PDF/DOCX/TXT
Algoritma 2: deteksi seksi (title, abstract, method, results, novelty)
"""
import re
from io import BytesIO


def extract_text(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    """Extract text from PDF/DOCX/TXT. Returns (text, metadata)."""
    name_lower = filename.lower()
    metadata = {"filename": filename, "size_kb": len(file_bytes) / 1024,
                "pages": 0, "format": "unknown"}

    try:
        if name_lower.endswith(".pdf"):
            return _extract_pdf(file_bytes, metadata)
        elif name_lower.endswith((".docx", ".doc")):
            return _extract_docx(file_bytes, metadata)
        elif name_lower.endswith(".txt"):
            metadata["format"] = "txt"
            text = file_bytes.decode("utf-8", errors="ignore")
            return text, metadata
        else:
            return "", metadata
    except Exception as e:
        metadata["error"] = str(e)
        return "", metadata


def _extract_pdf(file_bytes: bytes, metadata: dict) -> tuple[str, dict]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[pypdf belum terinstall]", metadata
    reader = PdfReader(BytesIO(file_bytes))
    metadata["pages"] = len(reader.pages)
    metadata["format"] = "pdf"
    text_parts = []
    for i, page in enumerate(reader.pages):
        try:
            text_parts.append(f"\n[PAGE {i+1}]\n" + (page.extract_text() or ""))
        except Exception:
            pass
    return "\n".join(text_parts), metadata


def _extract_docx(file_bytes: bytes, metadata: dict) -> tuple[str, dict]:
    try:
        from docx import Document
    except ImportError:
        return "[python-docx belum terinstall]", metadata
    doc = Document(BytesIO(file_bytes))
    metadata["format"] = "docx"
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Extract tables too
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                paragraphs.append("[TABLE] " + row_text)

    return "\n".join(paragraphs), metadata


def detect_sections(text: str) -> dict:
    """Detect IMRAD-like sections from text. Returns dict with section content."""
    if not text or len(text) < 50:
        return {}

    sections = {}
    lines = text.split("\n")

    # Title — first non-empty line under 200 chars, or first heading
    for line in lines:
        line_strip = line.strip()
        if 10 < len(line_strip) < 200 and not line_strip.startswith("[PAGE"):
            sections["title"] = line_strip
            break

    # Section patterns (case-insensitive)
    section_patterns = {
        "abstract":     r"(?i)abstract\b|\babstrak\b",
        "keywords":     r"(?i)keywords?\b|\bkata kunci\b",
        "introduction": r"(?i)\bintroduction\b|\bpendahuluan\b|\bbackground\b",
        "method":       r"(?i)\bmethod(s|ology)?\b|\bmetode\b|\bmetodologi\b",
        "result":       r"(?i)\bresult(s)?\b|\bhasil\b|\bfinding(s)?\b",
        "discussion":   r"(?i)\bdiscussion\b|\bdiskusi\b|\bpembahasan\b",
        "conclusion":   r"(?i)\bconclusion(s)?\b|\bkesimpulan\b",
        "novelty":      r"(?i)\bnovelty\b|\bkebaruan\b|novel|first time|propose",
    }

    for sec_name, pattern in section_patterns.items():
        match = re.search(pattern, text)
        if match:
            start = match.start()
            # Take ~1500 chars after the heading
            content = text[start:start + 1500].strip()
            sections[sec_name] = content

    sections["_full_text"] = text
    sections["_total_chars"] = len(text)
    return sections
